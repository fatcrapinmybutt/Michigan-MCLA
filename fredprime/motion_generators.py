"""Motion document generators for Michigan court practice.

Generates DOCX-format motions. Requires python-docx: pip install python-docx
"""

from __future__ import annotations

import os

try:
    from docx import Document
except ImportError:
    Document = None

DEFAULT_OUTPUT_DIR = os.path.join("output", "motions")


def _ensure_docx() -> None:
    if Document is None:
        raise ImportError("python-docx is required: pip install python-docx")


def generate_emergency_injunction(
    points: list[str] | None = None,
    output_dir: str = DEFAULT_OUTPUT_DIR,
    filename: str = "emergency_injunction_motion.docx",
) -> str:
    """Generate an emergency motion for preliminary injunction."""
    _ensure_docx()
    os.makedirs(output_dir, exist_ok=True)
    doc = Document()
    doc.add_heading("Emergency Motion for Preliminary Injunction", 0)
    for p in (points or []):
        doc.add_paragraph(p, style="List Bullet")
    out_path = os.path.join(output_dir, filename)
    doc.save(out_path)
    return out_path


def generate_protective_order(
    relief: str = "",
    authority: str = "MCR 2.302(C)",
    output_dir: str = DEFAULT_OUTPUT_DIR,
    filename: str = "motion_protective_order.docx",
) -> str:
    """Generate a motion for protective order."""
    _ensure_docx()
    os.makedirs(output_dir, exist_ok=True)
    doc = Document()
    doc.add_heading("Emergency Motion for Protective Order", 0)
    if relief:
        doc.add_paragraph(f"Relief requested: {relief}")
    doc.add_paragraph(f"Authority: {authority}")
    out_path = os.path.join(output_dir, filename)
    doc.save(out_path)
    return out_path
