"""Federal complaint document generator.

Generates DOCX-format federal complaints for civil rights actions.
Requires python-docx: pip install python-docx
"""

from __future__ import annotations

import os
from pathlib import Path

try:
    from docx import Document
except ImportError:
    Document = None

DEFAULT_OUTPUT_DIR = os.path.join("output", "federal")


def generate_complaint(
    court: str = "U.S. District Court",
    plaintiff: str = "",
    defendants: list[str] | None = None,
    counts: list[str] | None = None,
    allegations: list[str] | None = None,
    output_dir: str = DEFAULT_OUTPUT_DIR,
    filename: str = "federal_complaint.docx",
) -> str:
    """Generate a federal complaint document.

    Args:
        court: Court name for the caption.
        plaintiff: Plaintiff name.
        defendants: List of defendant names.
        counts: List of legal counts.
        allegations: List of key allegations.
        output_dir: Directory to save the output.
        filename: Output filename.

    Returns:
        Path to the generated document.
    """
    if Document is None:
        raise ImportError("python-docx is required: pip install python-docx")

    os.makedirs(output_dir, exist_ok=True)
    doc = Document()
    doc.add_heading("Federal Complaint", 0)
    doc.add_paragraph(f"Court: {court}")
    doc.add_paragraph(f"Plaintiff: {plaintiff}")

    if defendants:
        doc.add_paragraph("Defendants: " + ", ".join(defendants))

    if counts:
        doc.add_heading("Counts", level=1)
        for c in counts:
            doc.add_paragraph(c, style="List Number")

    if allegations:
        doc.add_heading("Key Allegations", level=1)
        for a in allegations:
            doc.add_paragraph(a, style="List Bullet")

    out_path = os.path.join(output_dir, filename)
    doc.save(out_path)
    return out_path
